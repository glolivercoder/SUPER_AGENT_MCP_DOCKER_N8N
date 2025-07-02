#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Advanced RAG System
------------------
Sistema RAG avançado para documentações técnicas, integração com OpenRouter
e processamento rápido de documentos PDF, texto, URLs e documentações.

Autor: [Seu Nome]
Data: 01/07/2025
Versão: 0.1.0
"""

import os
import json
import logging
import requests
import asyncio
from datetime import datetime
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import hashlib

# Processamento de documentos
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import nltk
    from nltk.tokenize import word_tokenize, sent_tokenize
    from nltk.corpus import stopwords
    NLTK_AVAILABLE = True
except ImportError:
    NLTK_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

logger = logging.getLogger("ADVANCED_RAG")

class DocumentProcessor:
    """Processador de diferentes tipos de documentos"""
    
    def __init__(self):
        self.logger = logger
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.docx': self._process_docx,
            '.html': self._process_html,
            '.json': self._process_json
        }
    
    def process_file(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo baseado na extensão"""
        try:
            extension = file_path.suffix.lower()
            
            if extension not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Formato {extension} não suportado",
                    "content": "",
                    "metadata": {}
                }
            
            processor = self.supported_formats[extension]
            return processor(file_path)
            
        except Exception as e:
            self.logger.error(f"Erro ao processar arquivo {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "metadata": {}
            }
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo PDF"""
        if not PDF_AVAILABLE:
            return {"success": False, "error": "PyPDF2 não disponível"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    content += f"\n--- Página {page_num + 1} ---\n"
                    content += page.extract_text()
                
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "title": getattr(pdf_reader.metadata, 'title', ''),
                    "author": getattr(pdf_reader.metadata, 'author', ''),
                    "format": "PDF"
                }
                
                return {
                    "success": True,
                    "content": content,
                    "metadata": metadata
                }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo de texto"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "lines": len(content.split('\n')),
                "words": len(content.split()),
                "format": "TEXT"
            }
            
            return {
                "success": True,
                "content": content,
                "metadata": metadata
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _process_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo Markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Converter para HTML se markdown disponível
            if MARKDOWN_AVAILABLE:
                html_content = markdown.markdown(md_content)
            else:
                html_content = md_content
            
            metadata = {
                "format": "MARKDOWN",
                "lines": len(md_content.split('\n')),
                "headers": len([line for line in md_content.split('\n') if line.startswith('#')])
            }
            
            return {
                "success": True,
                "content": md_content,
                "html_content": html_content,
                "metadata": metadata
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo DOCX"""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx não disponível"}
        
        try:
            doc = DocxDocument(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "format": "DOCX"
            }
            
            return {
                "success": True,
                "content": content,
                "metadata": metadata
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _process_html(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo HTML"""
        if not BS4_AVAILABLE:
            return {"success": False, "error": "BeautifulSoup não disponível"}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            metadata = {
                "title": soup.title.string if soup.title else "",
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img')),
                "format": "HTML"
            }
            
            return {
                "success": True,
                "content": text_content,
                "html_content": html_content,
                "metadata": metadata
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Processa arquivo JSON"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            
            content = json.dumps(json_data, indent=2, ensure_ascii=False)
            
            metadata = {
                "keys": len(json_data.keys()) if isinstance(json_data, dict) else 0,
                "format": "JSON"
            }
            
            return {
                "success": True,
                "content": content,
                "json_data": json_data,
                "metadata": metadata
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

class OpenRouterClient:
    """Cliente para integração com OpenRouter API"""
    
    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv("OPENROUTER_API_KEY")
        self.base_url = "https://openrouter.ai/api/v1"
        self.logger = logger
        self.models_cache = None
        
    def get_models(self) -> Dict[str, Any]:
        """Obtém lista de modelos disponíveis na OpenRouter"""
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            response = requests.get(f"{self.base_url}/models", headers=headers)
            
            if response.status_code == 200:
                models_data = response.json()
                self.models_cache = models_data
                
                # Organizar modelos por empresa
                models_by_company = {}
                free_models = []
                
                for model in models_data.get("data", []):
                    company = model.get("owned_by", "Unknown")
                    model_id = model.get("id", "")
                    
                    if company not in models_by_company:
                        models_by_company[company] = []
                    
                    models_by_company[company].append(model)
                    
                    # Verificar se é gratuito
                    pricing = model.get("pricing", {})
                    prompt_price = float(pricing.get("prompt", "0"))
                    completion_price = float(pricing.get("completion", "0"))
                    
                    if prompt_price == 0 and completion_price == 0:
                        free_models.append(model)
                
                return {
                    "success": True,
                    "models": models_data.get("data", []),
                    "models_by_company": models_by_company,
                    "free_models": free_models,
                    "total_models": len(models_data.get("data", []))
                }
            else:
                return {
                    "success": False,
                    "error": f"Erro {response.status_code}: {response.text}"
                }
        except Exception as e:
            self.logger.error(f"Erro ao obter modelos OpenRouter: {e}")
            return {"success": False, "error": str(e)}
    
    def query_model(self, model_id: str, prompt: str, context: str = "") -> Dict[str, Any]:
        """Faz query para um modelo específico"""
        try:
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://super-agent-mcp.local",
                "X-Title": "SUPER_AGENT_MCP_DOCKER_N8N"
            }
            
            messages = []
            if context:
                messages.append({
                    "role": "system",
                    "content": f"Contexto: {context}"
                })
            
            messages.append({
                "role": "user", 
                "content": prompt
            })
            
            data = {
                "model": model_id,
                "messages": messages,
                "max_tokens": 2000,
                "temperature": 0.7
            }
            
            response = requests.post(
                f"{self.base_url}/chat/completions", 
                headers=headers, 
                json=data
            )
            
            if response.status_code == 200:
                result = response.json()
                return {
                    "success": True,
                    "response": result.get("choices", [{}])[0].get("message", {}).get("content", ""),
                    "model_used": model_id,
                    "usage": result.get("usage", {})
                }
            else:
                return {
                    "success": False,
                    "error": f"Erro {response.status_code}: {response.text}"
                }
        except Exception as e:
            self.logger.error(f"Erro ao consultar modelo {model_id}: {e}")
            return {"success": False, "error": str(e)}

class AdvancedRAGSystem:
    """Sistema RAG avançado com documentações técnicas e OpenRouter"""
    
    def __init__(self, data_dir: str = "data", config_dir: str = "config"):
        self.data_dir = Path(data_dir)
        self.config_dir = Path(config_dir)
        self.logger = logger
        
        # Criar diretórios
        os.makedirs(self.data_dir, exist_ok=True)
        os.makedirs(self.data_dir / "docs", exist_ok=True)
        os.makedirs(self.data_dir / "cache", exist_ok=True)
        
        # Inicializar componentes
        self.document_processor = DocumentProcessor()
        self.openrouter_client = OpenRouterClient()
        
        # Base de conhecimento
        self.knowledge_base = {
            "documents": {},
            "documentation_index": {},
            "url_cache": {},
            "last_updated": None
        }
        
        # Documentações técnicas pré-configuradas
        self.tech_docs = {
            "git": {
                "name": "Git Documentation",
                "url": "https://git-scm.com/docs",
                "topics": ["commands", "workflow", "branching", "merging"]
            },
            "docker": {
                "name": "Docker Documentation", 
                "url": "https://docs.docker.com",
                "topics": ["containers", "images", "compose", "swarm"]
            },
            "n8n": {
                "name": "N8N Documentation",
                "url": "https://docs.n8n.io",
                "topics": ["workflows", "nodes", "automation", "self-hosting"]
            },
            "mcp": {
                "name": "MCP Servers Documentation",
                "url": "https://modelcontextprotocol.io/docs",
                "topics": ["protocol", "servers", "clients", "implementation"]
            },
            "kubernetes": {
                "name": "Kubernetes Documentation",
                "url": "https://kubernetes.io/docs",
                "topics": ["pods", "services", "deployment", "ingress"]
            },
            "cloudflare": {
                "name": "Cloudflare Workers",
                "url": "https://developers.cloudflare.com/workers",
                "topics": ["workers", "kv", "durable-objects", "pages"]
            },
            "digitalocean": {
                "name": "DigitalOcean Documentation",
                "url": "https://docs.digitalocean.com",
                "topics": ["droplets", "kubernetes", "app-platform", "spaces"]
            },
            "oracle": {
                "name": "Oracle Cloud Free Tier",
                "url": "https://docs.oracle.com/en-us/iaas/Content/FreeTier/freetier_topic-Landing.htm",
                "topics": ["compute", "storage", "networking", "database"]
            }
        }
        
        self._load_knowledge_base()
    
    def _load_knowledge_base(self):
        """Carrega base de conhecimento do arquivo"""
        kb_file = self.data_dir / "knowledge_base.json"
        
        if kb_file.exists():
            try:
                with open(kb_file, 'r', encoding='utf-8') as f:
                    self.knowledge_base = json.load(f)
                self.logger.info("Base de conhecimento carregada")
            except Exception as e:
                self.logger.error(f"Erro ao carregar base de conhecimento: {e}")
    
    def _save_knowledge_base(self):
        """Salva base de conhecimento no arquivo"""
        kb_file = self.data_dir / "knowledge_base.json"
        
        try:
            with open(kb_file, 'w', encoding='utf-8') as f:
                json.dump(self.knowledge_base, f, indent=2, ensure_ascii=False)
            self.logger.info("Base de conhecimento salva")
        except Exception as e:
            self.logger.error(f"Erro ao salvar base de conhecimento: {e}")
    
    def add_document(self, file_path: Union[str, Path], category: str = "general") -> Dict[str, Any]:
        """Adiciona documento à base de conhecimento"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {"success": False, "error": "Arquivo não encontrado"}
            
            # Processar documento
            processed = self.document_processor.process_file(file_path)
            
            if not processed["success"]:
                return processed
            
            # Gerar ID único
            file_hash = hashlib.md5(file_path.name.encode()).hexdigest()[:8]
            doc_id = f"{category}_{file_hash}_{int(datetime.now().timestamp())}"
            
            # Adicionar à base de conhecimento
            self.knowledge_base["documents"][doc_id] = {
                "id": doc_id,
                "filename": file_path.name,
                "filepath": str(file_path),
                "category": category,
                "content": processed["content"],
                "metadata": processed["metadata"],
                "added_at": datetime.now().isoformat(),
                "hash": file_hash
            }
            
            # Atualizar índice
            self._update_documentation_index(doc_id, processed["content"], category)
            
            # Salvar
            self._save_knowledge_base()
            
            return {
                "success": True,
                "document_id": doc_id,
                "message": f"Documento {file_path.name} adicionado com sucesso"
            }
            
        except Exception as e:
            self.logger.error(f"Erro ao adicionar documento: {e}")
            return {"success": False, "error": str(e)}
    
    def add_url_content(self, url: str, category: str = "web") -> Dict[str, Any]:
        """Adiciona conteúdo de URL à base de conhecimento"""
        try:
            # Verificar cache
            url_hash = hashlib.md5(url.encode()).hexdigest()
            
            if url_hash in self.knowledge_base.get("url_cache", {}):
                cached = self.knowledge_base["url_cache"][url_hash]
                # Verificar se cache ainda é válido (24 horas)
                cache_time = datetime.fromisoformat(cached["cached_at"])
                if (datetime.now() - cache_time).hours < 24:
                    return {"success": True, "message": "Conteúdo já em cache", "cached": True}
            
            # Baixar conteúdo
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # Processar HTML
            if BS4_AVAILABLE:
                soup = BeautifulSoup(response.text, 'html.parser')
                content = soup.get_text()
                title = soup.title.string if soup.title else url
            else:
                content = response.text
                title = url
            
            # Gerar ID
            doc_id = f"{category}_url_{url_hash}"
            
            # Adicionar à base
            self.knowledge_base["documents"][doc_id] = {
                "id": doc_id,
                "url": url,
                "title": title,
                "category": category,
                "content": content,
                "metadata": {"format": "URL", "length": len(content)},
                "added_at": datetime.now().isoformat()
            }
            
            # Cache URL
            self.knowledge_base["url_cache"][url_hash] = {
                "url": url,
                "doc_id": doc_id,
                "cached_at": datetime.now().isoformat()
            }
            
            # Atualizar índice
            self._update_documentation_index(doc_id, content, category)
            
            # Salvar
            self._save_knowledge_base()
            
            return {
                "success": True,
                "document_id": doc_id,
                "message": f"Conteúdo de {url} adicionado com sucesso"
            }
            
        except Exception as e:
            self.logger.error(f"Erro ao adicionar URL {url}: {e}")
            return {"success": False, "error": str(e)}
    
    def _update_documentation_index(self, doc_id: str, content: str, category: str):
        """Atualiza índice de documentação para busca rápida"""
        if "documentation_index" not in self.knowledge_base:
            self.knowledge_base["documentation_index"] = {}
        
        # Extrair palavras-chave
        if NLTK_AVAILABLE:
            try:
                words = word_tokenize(content.lower())
                stop_words = set(stopwords.words('english') + stopwords.words('portuguese'))
                keywords = [word for word in words if word.isalnum() and word not in stop_words]
            except:
                keywords = content.lower().split()
        else:
            keywords = content.lower().split()
        
        # Adicionar ao índice
        for keyword in set(keywords[:100]):  # Limitar a 100 palavras-chave
            if keyword not in self.knowledge_base["documentation_index"]:
                self.knowledge_base["documentation_index"][keyword] = []
            
            if doc_id not in self.knowledge_base["documentation_index"][keyword]:
                self.knowledge_base["documentation_index"][keyword].append(doc_id)
    
    def search_knowledge(self, query: str, category: str = None, limit: int = 5) -> Dict[str, Any]:
        """Busca rápida na base de conhecimento"""
        try:
            query_lower = query.lower()
            results = []
            
            # Buscar por palavras-chave no índice
            query_words = query_lower.split()
            matching_docs = set()
            
            for word in query_words:
                if word in self.knowledge_base.get("documentation_index", {}):
                    matching_docs.update(self.knowledge_base["documentation_index"][word])
            
            # Filtrar por categoria se especificada
            for doc_id in matching_docs:
                if doc_id in self.knowledge_base["documents"]:
                    doc = self.knowledge_base["documents"][doc_id]
                    
                    if category and doc.get("category") != category:
                        continue
                    
                    # Calcular relevância simples
                    content_lower = doc["content"].lower()
                    relevance = sum(1 for word in query_words if word in content_lower)
                    
                    results.append({
                        "document_id": doc_id,
                        "title": doc.get("title", doc.get("filename", "Sem título")),
                        "category": doc.get("category", "general"),
                        "relevance": relevance,
                        "snippet": content_lower[:200] + "..." if len(content_lower) > 200 else content_lower,
                        "metadata": doc.get("metadata", {})
                    })
            
            # Ordenar por relevância
            results.sort(key=lambda x: x["relevance"], reverse=True)
            
            return {
                "success": True,
                "query": query,
                "results": results[:limit],
                "total_found": len(results)
            }
            
        except Exception as e:
            self.logger.error(f"Erro na busca: {e}")
            return {"success": False, "error": str(e)}
    
    def get_openrouter_models(self) -> Dict[str, Any]:
        """Obtém modelos da OpenRouter organizados"""
        return self.openrouter_client.get_models()
    
    def query_with_context(self, question: str, model_id: str, context_docs: List[str] = None) -> Dict[str, Any]:
        """Faz pergunta com contexto da base de conhecimento"""
        try:
            # Buscar contexto se não fornecido
            if not context_docs:
                search_result = self.search_knowledge(question, limit=3)
                if search_result["success"]:
                    context_docs = [doc["document_id"] for doc in search_result["results"]]
            
            # Construir contexto
            context_content = ""
            for doc_id in context_docs:
                if doc_id in self.knowledge_base["documents"]:
                    doc = self.knowledge_base["documents"][doc_id]
                    context_content += f"\n--- {doc.get('title', doc_id)} ---\n"
                    context_content += doc["content"][:1000]  # Limitar contexto
                    context_content += "\n"
            
            # Fazer query
            result = self.openrouter_client.query_model(model_id, question, context_content)
            
            return {
                "success": result["success"],
                "question": question,
                "answer": result.get("response", ""),
                "model_used": model_id,
                "context_documents": context_docs,
                "usage": result.get("usage", {}),
                "error": result.get("error")
            }
            
        except Exception as e:
            self.logger.error(f"Erro na query com contexto: {e}")
            return {"success": False, "error": str(e)}
    
    def get_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas da base de conhecimento"""
        docs = self.knowledge_base.get("documents", {})
        
        # Contar por categoria
        categories = {}
        total_size = 0
        
        for doc in docs.values():
            category = doc.get("category", "general")
            categories[category] = categories.get(category, 0) + 1
            total_size += len(doc.get("content", ""))
        
        return {
            "total_documents": len(docs),
            "categories": categories,
            "total_content_size": total_size,
            "index_keywords": len(self.knowledge_base.get("documentation_index", {})),
            "cached_urls": len(self.knowledge_base.get("url_cache", {})),
            "last_updated": self.knowledge_base.get("last_updated")
        }

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    rag = AdvancedRAGSystem()
    models = rag.get_openrouter_models()
    print(f"Sistema RAG inicializado") 